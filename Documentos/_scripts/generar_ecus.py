"""Genera el documento de Especificaciones de Casos de Uso (ECUs) en Word."""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RAIZ = Path(__file__).resolve().parents[1]
SALIDA = RAIZ / "Tecnico" / "ECUs-MedScribe.docx"
SALIDA.parent.mkdir(parents=True, exist_ok=True)

AZUL = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()
for s in doc.sections:
    s.page_height = Cm(29.7)
    s.page_width = Cm(21.0)
    s.top_margin = Cm(3); s.bottom_margin = Cm(3)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Arial"; style.font.size = Pt(11)


def titulo(txt, nivel=1, color=AZUL, centrar=False):
    p = doc.add_paragraph()
    if centrar: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.bold = True; r.font.name = "Arial"; r.font.color.rgb = color
    r.font.size = {0: Pt(20), 1: Pt(16), 2: Pt(13), 3: Pt(11)}[nivel]


def parrafo(txt, justificar=True):
    p = doc.add_paragraph()
    if justificar: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(txt); r.font.name = "Arial"; r.font.size = Pt(11)


def ecu_tabla(datos):
    """datos = lista de [campo, valor]"""
    t = doc.add_table(rows=len(datos), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (campo, valor) in enumerate(datos):
        c1 = t.rows[i].cells[0]; c2 = t.rows[i].cells[1]
        c1.text = ""; c2.text = ""
        p1 = c1.paragraphs[0]; r1 = p1.add_run(campo)
        r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Arial"
        r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = c1._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "1F4E79"); tcPr.append(shd)
        c1.width = Cm(4)
        p2 = c2.paragraphs[0]
        if isinstance(valor, list):
            for j, item in enumerate(valor):
                if j > 0: p2 = c2.add_paragraph()
                r2 = p2.add_run(item)
                r2.font.size = Pt(10); r2.font.name = "Arial"
        else:
            r2 = p2.add_run(str(valor))
            r2.font.size = Pt(10); r2.font.name = "Arial"
        c2.width = Cm(11)
    doc.add_paragraph()


# ===========
# CARATULA
# ===========
for _ in range(4): doc.add_paragraph()
titulo("INSTITUTO CIBERTEC", nivel=1, centrar=True)
titulo("Tecnologias de la Informacion", nivel=2, centrar=True)
for _ in range(2): doc.add_paragraph()
titulo("ESPECIFICACIONES DE CASOS DE USO (ECUs)", nivel=1, centrar=True)
for _ in range(2): doc.add_paragraph()
titulo("MEDSCRIBE AI", nivel=0, centrar=True)
parrafo("Sistema de Documentacion Medica Automatizada con IA")
for _ in range(4): doc.add_paragraph()
parrafo("Curso: Desarrollo de Servicios Web I (4695)")
parrafo("Docente: Ing. Antonio Valdivia")
parrafo("Equipo: Roberto La Rosa, Jason Davila, Luis Curi, Edward Escobedo")
parrafo("Fecha: Mayo 2026")
doc.add_page_break()

# ===========
# INTRODUCCION
# ===========
titulo("1. INTRODUCCION", nivel=1)
parrafo(
    "Este documento detalla la especificacion textual de los casos de uso del "
    "sistema MedScribe AI. Cada ECU describe quien dispara la accion, los pasos "
    "del flujo principal, las variantes (flujos alternativos) y las condiciones "
    "que deben cumplirse antes y despues de la operacion."
)
parrafo(
    "El sistema reconoce dos actores: el Medico (quien atiende a los pacientes) "
    "y el Administrador (quien ademas gestiona la clinica). Los casos de uso "
    "estan agrupados por modulo funcional."
)

titulo("2. LISTADO DE CASOS DE USO", nivel=1)
parrafo("Se identificaron 14 casos de uso, distribuidos asi:")

t = doc.add_table(rows=15, cols=4)
t.style = "Light Grid Accent 1"
headers = ["Codigo", "Caso de Uso", "Modulo", "Actor(es)"]
for i, h in enumerate(headers):
    c = t.rows[0].cells[i]; c.text = ""
    r = c.paragraphs[0].add_run(h)
    r.bold = True; r.font.size = Pt(10); r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tcPr = c._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "1F4E79"); tcPr.append(shd)

casos = [
    ("CU-01", "Iniciar Sesion", "Autenticacion", "Medico, Admin"),
    ("CU-02", "Cerrar Sesion", "Autenticacion", "Medico, Admin"),
    ("CU-03", "Gestionar Pacientes", "Pacientes", "Medico, Admin"),
    ("CU-04", "Crear Consulta desde Audio", "Consultas", "Medico, Admin"),
    ("CU-05", "Transcribir y Estructurar Nota", "Consultas (incluido)", "Sistema"),
    ("CU-06", "Editar Nota Clinica", "Consultas", "Medico, Admin"),
    ("CU-07", "Generar PDF y Word", "Documentos", "Medico, Admin"),
    ("CU-08", "Ver Historial de Consultas", "Consultas", "Medico, Admin"),
    ("CU-09", "Aprobar Documento", "Documentos", "Admin"),
    ("CU-10", "Gestionar Usuarios de la Clinica", "Usuarios", "Admin"),
    ("CU-11", "Gestionar Roles y Permisos", "Roles", "Admin"),
    ("CU-12", "Configurar Plantillas de Documentos", "Plantillas", "Admin"),
    ("CU-13", "Editar Perfil", "Perfil", "Medico, Admin"),
    ("CU-14", "Cambiar Contrasena", "Perfil", "Medico, Admin"),
]
for i, (cod, cu, mod, act) in enumerate(casos, 1):
    for j, val in enumerate([cod, cu, mod, act]):
        c = t.rows[i].cells[j]; c.text = ""
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(10); r.font.name = "Arial"
doc.add_paragraph()
doc.add_page_break()

# ===========
# ECUs DETALLADAS
# ===========
titulo("3. ESPECIFICACIONES DETALLADAS", nivel=1)

ecus_data = [
    {
        "cod": "CU-01", "nombre": "Iniciar Sesion",
        "campos": [
            ("Codigo", "CU-01"),
            ("Nombre", "Iniciar Sesion"),
            ("Actor", "Medico, Administrador"),
            ("Modulo", "Autenticacion"),
            ("Descripcion", "Permite al usuario acceder al sistema mediante su correo y contrasena."),
            ("Precondicion", "El usuario debe estar registrado y activo en una clinica."),
            ("Flujo principal", [
                "1. El usuario ingresa a la pagina de inicio de sesion.",
                "2. Escribe su correo electronico y su contrasena.",
                "3. Hace click en el boton 'Iniciar sesion'.",
                "4. El sistema valida el formato del correo y la contrasena.",
                "5. El sistema envia las credenciales al servidor.",
                "6. El servidor ejecuta el stored procedure usp_Usuarios_ValidarCredenciales.",
                "7. Si las credenciales son correctas, el servidor genera un token JWT.",
                "8. El sistema guarda el token en localStorage.",
                "9. El sistema redirige al usuario al Panel principal.",
            ]),
            ("Flujos alternativos", [
                "5a. El correo no existe: el sistema muestra 'Correo o contrasena incorrectos'.",
                "5b. La contrasena es incorrecta: igual mensaje generico (por seguridad).",
                "5c. El usuario esta inactivo: el sistema muestra 'Cuenta desactivada, contacte al administrador'.",
            ]),
            ("Postcondicion", "El usuario queda autenticado y se le otorgan los permisos de su rol."),
        ],
    },
    {
        "cod": "CU-04", "nombre": "Crear Consulta desde Audio",
        "campos": [
            ("Codigo", "CU-04"),
            ("Nombre", "Crear Consulta desde Audio"),
            ("Actor", "Medico, Administrador"),
            ("Modulo", "Consultas"),
            ("Descripcion", "El medico graba o sube un audio de la consulta y el sistema genera automaticamente la nota clinica estructurada."),
            ("Precondicion", [
                "El usuario debe estar autenticado.",
                "Debe existir al menos un paciente registrado.",
                "El usuario debe tener permiso 'consultas.crear'.",
            ]),
            ("Flujo principal", [
                "1. El usuario hace click en 'Nueva Consulta'.",
                "2. Selecciona al paciente de la lista.",
                "3. Selecciona la plantilla de nota clinica (SOAP, Receta, etc.).",
                "4. Graba el audio con el microfono o sube un archivo .wav/.mp3.",
                "5. Hace click en 'Procesar consulta'.",
                "6. El sistema invoca el caso de uso CU-05 (Transcribir y Estructurar).",
                "7. El sistema muestra la nota clinica generada en el editor.",
                "8. El usuario revisa y, si es necesario, edita el contenido.",
                "9. Hace click en 'Guardar consulta'.",
                "10. El sistema guarda la consulta y los valores de seccion en SQL Server.",
                "11. El sistema invoca CU-07 para generar PDF y Word.",
                "12. Muestra confirmacion con enlaces de descarga.",
            ]),
            ("Flujos alternativos", [
                "4a. Audio invalido (formato no soportado): muestra error y solicita otro archivo.",
                "5a. El servicio IA no responde en 90 s: muestra mensaje 'Reintente en unos segundos'.",
                "8a. La nota generada esta vacia o incompleta: el usuario puede repetir la transcripcion.",
            ]),
            ("Postcondicion", "Queda registrada una nueva consulta con su nota clinica y los documentos PDF y Word generados."),
        ],
    },
    {
        "cod": "CU-05", "nombre": "Transcribir y Estructurar Nota Clinica",
        "campos": [
            ("Codigo", "CU-05"),
            ("Nombre", "Transcribir y Estructurar Nota Clinica"),
            ("Actor", "Sistema (caso de uso incluido por CU-04)"),
            ("Modulo", "Servicio IA"),
            ("Descripcion", "Pipeline interno que transforma un audio en una nota clinica estructurada."),
            ("Precondicion", "Existe un archivo de audio valido recibido del frontend."),
            ("Flujo principal", [
                "1. El gateway envia el audio al servicio IA en POST /api/ia/transcribir.",
                "2. Whisper convierte el audio en texto plano.",
                "3. El diarizador (Deepgram o Pyannote) etiqueta cada parrafo con su orador.",
                "4. El gateway invoca POST /api/ia/procesar con la transcripcion.",
                "5. El clasificador de intenciones identifica el tipo de consulta.",
                "6. El servicio RAG busca contexto medico en Qdrant.",
                "7. Llama 3.3 genera la nota clinica estructurada en JSON.",
                "8. El servicio IA devuelve el JSON al gateway.",
            ]),
            ("Flujos alternativos", [
                "2a. Whisper falla por audio corrupto: el sistema reintenta una vez y, si vuelve a fallar, devuelve error 422.",
                "7a. Llama no responde dentro de 120 s: el circuit breaker se abre y devuelve un mensaje al usuario.",
                "7b. La respuesta de Llama es invalida: el sistema reintenta con backoff exponencial hasta 3 veces.",
            ]),
            ("Postcondicion", "El gateway recibe la nota estructurada lista para guardar."),
        ],
    },
    {
        "cod": "CU-07", "nombre": "Generar PDF y Word",
        "campos": [
            ("Codigo", "CU-07"),
            ("Nombre", "Generar PDF y Word"),
            ("Actor", "Sistema (a peticion del usuario)"),
            ("Modulo", "Documentos"),
            ("Descripcion", "Convierte la nota clinica estructurada en archivos PDF y Word listos para imprimir o descargar."),
            ("Precondicion", "Existe una consulta guardada con su nota clinica completa."),
            ("Flujo principal", [
                "1. El gateway invoca POST /api/ia/generar-pdf con el id de la consulta.",
                "2. El servicio IA arma el documento PDF con ReportLab usando la plantilla configurada.",
                "3. Aplica colores, logo de la clinica y formato profesional.",
                "4. Guarda el PDF en /documentos-salida con un nombre unico.",
                "5. Guarda metadata (id, ruta, fecha) en un sidecar JSON.",
                "6. Repite el proceso para Word con python-docx.",
                "7. Devuelve las rutas al gateway.",
                "8. El gateway registra los documentos en SQL Server.",
            ]),
            ("Flujos alternativos", [
                "4a. Sin espacio en disco: el sistema devuelve error 507 y notifica al admin.",
            ]),
            ("Postcondicion", "Existen dos archivos (PDF y Word) descargables y registrados en BD."),
        ],
    },
    {
        "cod": "CU-10", "nombre": "Gestionar Usuarios de la Clinica",
        "campos": [
            ("Codigo", "CU-10"),
            ("Nombre", "Gestionar Usuarios de la Clinica"),
            ("Actor", "Administrador"),
            ("Modulo", "Usuarios"),
            ("Descripcion", "Permite al administrador crear, editar, activar/desactivar y eliminar usuarios de su clinica."),
            ("Precondicion", "El usuario debe estar autenticado como Administrador."),
            ("Flujo principal", [
                "1. El admin entra al modulo 'Usuarios de la Clinica'.",
                "2. Ve la lista de usuarios actuales con su rol y estado.",
                "3. Hace click en 'Nuevo Usuario' o en el lapiz para editar uno existente.",
                "4. Llena el formulario (correo, nombres, rol, especialidad si es medico).",
                "5. Hace click en 'Guardar'.",
                "6. El sistema valida los datos y, si todo esta bien, ejecuta el SP correspondiente.",
                "7. Muestra notificacion de exito y refresca la lista.",
            ]),
            ("Flujos alternativos", [
                "4a. Correo ya existe en la clinica: muestra 'El correo ya esta en uso'.",
                "6a. Validacion fallida: resalta los campos con error.",
            ]),
            ("Postcondicion", "El nuevo usuario o la modificacion queda registrado en la base de datos."),
        ],
    },
    {
        "cod": "CU-11", "nombre": "Gestionar Roles y Permisos",
        "campos": [
            ("Codigo", "CU-11"),
            ("Nombre", "Gestionar Roles y Permisos"),
            ("Actor", "Administrador"),
            ("Modulo", "Roles"),
            ("Descripcion", "Permite crear roles personalizados y asignar permisos granulares por modulo y accion."),
            ("Precondicion", "El usuario debe ser Administrador."),
            ("Flujo principal", [
                "1. El admin entra al modulo 'Roles'.",
                "2. Ve los roles existentes (Medico, Admin y los personalizados).",
                "3. Crea un nuevo rol o edita uno existente.",
                "4. En una matriz de permisos, marca o desmarca las casillas para cada modulo (pacientes, consultas, documentos, configuracion, usuarios, roles) y cada accion (ver, crear, editar, eliminar).",
                "5. Hace click en 'Guardar'.",
                "6. El sistema serializa los permisos como JSON y los guarda en la tabla RolesDeClinica.",
            ]),
            ("Flujos alternativos", [
                "5a. Intenta eliminar un rol con usuarios asignados: el sistema lo impide.",
            ]),
            ("Postcondicion", "El rol queda creado o actualizado y se aplica a los usuarios asignados."),
        ],
    },
]

for ecu in ecus_data:
    titulo(f"{ecu['cod']}: {ecu['nombre']}", nivel=2)
    ecu_tabla(ecu["campos"])

# ===========
# Resto resumido
# ===========
titulo("4. ECUs RESUMIDAS (CU-02, CU-03, CU-06, CU-08, CU-09, CU-12, CU-13, CU-14)", nivel=1)
parrafo(
    "Por brevedad, los siguientes casos de uso siguen el mismo patron descriptivo y se "
    "encuentran completamente implementados. Cada uno cuenta con su flujo principal en "
    "el codigo fuente, en los controladores correspondientes del Gateway C# y en las "
    "paginas React asociadas."
)

resumen_otros = [
    ("CU-02 Cerrar Sesion", "Borra el token JWT del localStorage y redirige a la pagina de login. Sin precondiciones especiales."),
    ("CU-03 Gestionar Pacientes", "CRUD completo de pacientes con validacion de DNI peruano (8 digitos), formato de telefono y correo. SP transaccional para alta y baja."),
    ("CU-06 Editar Nota Clinica", "El usuario puede modificar el contenido de cada seccion de la nota antes de aprobarla. Se guarda en la tabla ValoresDeSeccion."),
    ("CU-08 Ver Historial de Consultas", "Lista paginada con filtros por paciente, fecha y estado. Incluye busqueda por DNI."),
    ("CU-09 Aprobar Documento", "El admin marca el documento como APROBADO y queda inmutable. Se registra en AuditoriaDeConsultas."),
    ("CU-12 Configurar Plantillas", "Permite crear nuevas plantillas (SOAP, Historia Clinica, Receta) y definir las secciones que la IA debe rellenar."),
    ("CU-13 Editar Perfil", "El usuario actualiza sus datos personales (nombres, telefono, especialidad)."),
    ("CU-14 Cambiar Contrasena", "Requiere la contrasena actual + la nueva (con validacion de complejidad). Aplica BCrypt."),
]
for nombre, desc in resumen_otros:
    titulo(nombre, nivel=2)
    parrafo(desc)

doc.save(SALIDA)
print(f"OK: {SALIDA}")
print(f"Tamano: {SALIDA.stat().st_size // 1024} KB")
