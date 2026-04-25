"""Genera el Informe de Proyecto MedScribe en formato Word (.docx)."""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RAIZ = Path(__file__).resolve().parents[1]
DIAGRAMAS = RAIZ / "Tecnico" / "Diagramas"
SALIDA = RAIZ / "Negocio" / "Informe-Proyecto-MedScribe.docx"
SALIDA.parent.mkdir(parents=True, exist_ok=True)

AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x59, 0x59, 0x59)

doc = Document()

# Configuracion de pagina (segun rubrica: A4, margenes 3 sup/inf, 2.5 izq/der)
for seccion in doc.sections:
    seccion.page_height = Cm(29.7)
    seccion.page_width = Cm(21.0)
    seccion.top_margin = Cm(3)
    seccion.bottom_margin = Cm(3)
    seccion.left_margin = Cm(2.5)
    seccion.right_margin = Cm(2.5)

# Estilo base Arial 11
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)


def agregar_titulo(texto, nivel=1, color=AZUL):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(texto)
    run.bold = True
    run.font.name = "Arial"
    run.font.color.rgb = color
    if nivel == 0:
        run.font.size = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif nivel == 1:
        run.font.size = Pt(16)
    elif nivel == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    return p


def parrafo(texto, justificar=True, negrita=False, italica=False):
    p = doc.add_paragraph()
    if justificar:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texto)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = negrita
    run.italic = italica
    return p


def lista(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(11)


def insertar_imagen(nombre_png, titulo, fuente="Elaboracion propia", ancho_cm=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(DIAGRAMAS / nombre_png), width=Cm(ancho_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = cap.add_run(titulo + "\n")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = "Arial"
    r2 = cap.add_run("Fuente: " + fuente)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.name = "Arial"
    r2.font.color.rgb = GRIS


def tabla(encabezados, filas, anchos=None):
    t = doc.add_table(rows=1 + len(filas), cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(encabezados):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Sombrear celda
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)
    for ri, fila in enumerate(filas, start=1):
        for ci, valor in enumerate(fila):
            celda = t.rows[ri].cells[ci]
            celda.text = ""
            p = celda.paragraphs[0]
            run = p.add_run(str(valor))
            run.font.size = Pt(10)
            run.font.name = "Arial"
    if anchos:
        for fila in t.rows:
            for i, ancho in enumerate(anchos):
                fila.cells[i].width = Cm(ancho)
    doc.add_paragraph()


def salto_pagina():
    doc.add_page_break()


# =========================
# CARATULA
# =========================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INSTITUTO DE EDUCACION SUPERIOR CIBERTEC")
run.bold = True
run.font.size = Pt(14)
run.font.name = "Arial"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Tecnologias de la Informacion - Computacion e Informatica")
r.font.size = Pt(12)
r.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PROYECTO DE INVESTIGACION APLICADA")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = AZUL

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MEDSCRIBE AI")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = AZUL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema de Documentacion Medica Automatizada\ncon Inteligencia Artificial")
r.font.size = Pt(14)
r.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Curso: Desarrollo de Servicios Web I (4695)\n")
r.font.size = Pt(12)
r.bold = True
r = p.add_run("Docente: Ing. Antonio Valdivia\n")
r.font.size = Pt(12)
r = p.add_run("Ciclo: Quinto - Semestre 2026-I")
r.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run("Coordinador del grupo:\n")
r.bold = True
r.font.size = Pt(11)
r = p.add_run("    Jose Roberto La Rosa Ledezma  -  i202333980\n")
r.font.size = Pt(11)
r = p.add_run("\nIntegrantes del grupo:\n")
r.bold = True
r.font.size = Pt(11)
r = p.add_run(
    "    Jose Roberto La Rosa Ledezma  -  i202333980\n"
    "    Jason Davila Delgado  -  i202415540\n"
    "    Luis Joel Curi  -  i202417794\n"
    "    Edward Alexander Escobedo Murga  -  i201917851\n"
)
r.font.size = Pt(11)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Lima, Peru - Mayo 2026")
r.italic = True
r.font.size = Pt(11)

salto_pagina()

# =========================
# INDICE
# =========================
agregar_titulo("INDICE", nivel=1)
indice = [
    "1. Resumen",
    "2. Introduccion",
    "3. Diagnostico (Analisis SEPTE)",
    "    3.1 Variable Social",
    "    3.2 Variable Economica",
    "    3.3 Variable Tecnologica",
    "    3.4 Variable Politica",
    "4. Objetivos",
    "5. Justificacion del Proyecto",
    "6. Definicion y Alcance",
    "    6.1 Vision general",
    "    6.2 Arquitectura del sistema",
    "    6.3 Modulos funcionales",
    "    6.4 Tecnologias utilizadas",
    "    6.5 Diagramas RUP",
    "        6.5.1 Diagrama de Casos de Uso",
    "        6.5.2 Diagrama de Clases",
    "        6.5.3 Modelo Entidad-Relacion",
    "        6.5.4 Diagramas de Secuencia",
    "        6.5.5 Diagrama de Actividades",
    "        6.5.6 Diagrama de Componentes",
    "        6.5.7 Diagrama de Despliegue",
    "        6.5.8 Diagrama de Estados",
    "7. Productos y Entregables",
    "8. Conclusiones",
    "9. Recomendaciones",
    "10. Glosario",
    "11. Bibliografia",
    "12. Anexos",
]
for item in indice:
    p = doc.add_paragraph(item)
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.name = "Arial"

salto_pagina()

# =========================
# 1. RESUMEN
# =========================
agregar_titulo("1. RESUMEN", nivel=1)
parrafo(
    "MedScribe AI es una aplicacion web pensada para clinicas y consultorios medicos "
    "que automatiza la creacion de notas clinicas. El medico graba la consulta con su "
    "voz, el sistema escucha el audio, lo convierte en texto, identifica quien hablo "
    "(medico o paciente), entiende el contenido medico y arma una nota clinica "
    "estructurada lista para imprimir o guardar como PDF o Word."
)
parrafo(
    "El proyecto fue construido con cuatro tecnologias principales: una interfaz web "
    "en React, un backend en C# .NET 9 que actua como puerta de entrada, un servicio "
    "de inteligencia artificial en Python con FastAPI, y una base de datos SQL Server. "
    "Todo se ejecuta en contenedores Docker, de modo que cualquier persona puede "
    "instalarlo con un solo comando."
)
parrafo(
    "El alcance cubre el ciclo completo: desde el registro de la clinica y sus usuarios, "
    "la gestion de pacientes, la generacion automatica de notas clinicas a partir de "
    "audio, hasta la aprobacion y descarga de los documentos finales. Tambien incluye "
    "control de acceso por roles, multiples plantillas de documentos y un panel de "
    "estadisticas."
)

# =========================
# 2. INTRODUCCION
# =========================
agregar_titulo("2. INTRODUCCION", nivel=1)
parrafo(
    "En la atencion medica diaria, los profesionales pasan demasiado tiempo escribiendo "
    "notas en lugar de mirar al paciente. Estudios indican que un medico llega a "
    "dedicar entre el 35 % y el 50 % de su jornada a tareas administrativas y de "
    "documentacion. Esto reduce la calidad del contacto humano, aumenta el agotamiento "
    "(sindrome de burnout) y eleva la probabilidad de errores en la historia clinica."
)
parrafo(
    "MedScribe nace como respuesta a ese problema. La idea es sencilla: el medico habla "
    "como siempre lo hace y la computadora se encarga de escribir. Para lograrlo "
    "combinamos tres tecnologias de inteligencia artificial: transcripcion de voz a "
    "texto (Whisper), separacion de oradores (diarizacion con Deepgram o Pyannote) y "
    "modelos de lenguaje grande (Llama 3.3) que convierten el dialogo en una nota "
    "clinica estructurada."
)
parrafo(
    "Los objetivos del informe son: (a) mostrar el problema real que motiva el proyecto, "
    "(b) explicar como esta construida la solucion, (c) demostrar que la aplicacion "
    "funciona de extremo a extremo y (d) dejar evidencia de los entregables exigidos "
    "por el curso de Desarrollo de Servicios Web I."
)

salto_pagina()

# =========================
# 3. DIAGNOSTICO SEPTE
# =========================
agregar_titulo("3. DIAGNOSTICO (ANALISIS SEPTE)", nivel=1)
parrafo(
    "El analisis SEPTE permite ver el problema desde varios angulos. Hemos desarrollado "
    "cuatro variables (Social, Economica, Tecnologica y Politica) porque la rubrica del "
    "curso pide minimo tres incluyendo Tecnologica."
)

agregar_titulo("3.1 Variable Social", nivel=2)
parrafo(
    "En el Peru y en gran parte de America Latina existe una sobrecarga de los "
    "profesionales de salud. Segun el Colegio Medico del Peru (CMP, 2024), el pais "
    "cuenta con 13 medicos por cada 10 000 habitantes, cifra muy por debajo del "
    "estandar de 23 que recomienda la Organizacion Mundial de la Salud. Esto significa "
    "que cada medico atiende mas pacientes de lo deseable y dispone de menos tiempo "
    "para escuchar."
)
parrafo(
    "Una encuesta nacional del Instituto Nacional de Estadistica (INEI, 2023) sobre "
    "calidad de la atencion en EsSalud y MINSA encontro que el 63 % de los pacientes "
    "considera que su medico le dedica menos de 10 minutos por consulta. Tambien "
    "reportan que el medico pasa la mayor parte de ese tiempo escribiendo en la "
    "computadora en lugar de mirarlos a los ojos."
)
parrafo(
    "MedScribe ataca directamente esta variable: si la maquina escribe por el medico, "
    "el medico recupera ese tiempo para mirar y escuchar al paciente."
)

agregar_titulo("3.2 Variable Economica", nivel=2)
parrafo(
    "El costo de un sistema de historia clinica electronica importado (por ejemplo "
    "Epic, Cerner o Allscripts) supera los 100 000 dolares anuales por clinica, cifra "
    "fuera del alcance de consultorios pequenos. En Peru existen mas de 25 000 "
    "consultorios privados y solo el 18 % tiene un sistema digital propio (MINSA, 2023)."
)
parrafo(
    "Reducir el tiempo de documentacion en 30 minutos por dia significa, para una "
    "clinica con 10 medicos y honorarios promedio de S/ 80 por hora, un ahorro de "
    "aproximadamente S/ 96 000 anuales. MedScribe es ofrecido como servicio web con un "
    "plan basico al alcance de consultorios medianos y pequenos."
)

agregar_titulo("3.3 Variable Tecnologica", nivel=2)
parrafo(
    "Hasta hace tres anos, transcribir voz medica con buena precision requeria "
    "servidores costosos y modelos entrenados a la medida. Hoy, gracias a modelos como "
    "Whisper de OpenAI (open source) y APIs gratuitas como Groq que sirven Llama 3.3 "
    "con latencias de menos de un segundo, cualquier desarrollador puede integrar "
    "estas capacidades en una aplicacion web."
)
parrafo(
    "El stack que utilizamos (React, .NET 9, FastAPI, SQL Server, Qdrant, Docker) es "
    "moderno, mantenido por sus comunidades y compatible entre si. Esto permite que "
    "el sistema se ejecute en una laptop comun o en la nube, sin cambios en el codigo."
)

agregar_titulo("3.4 Variable Politica", nivel=2)
parrafo(
    "La Ley N.° 30024 que crea el Registro Nacional de Historias Clinicas Electronicas "
    "(RENHICE) y su reglamento (DS 009-2017-SA) obligan progresivamente a las "
    "instituciones de salud a digitalizar sus expedientes y a estandarizar el formato "
    "de las notas clinicas. La Ley N.° 29733 de Proteccion de Datos Personales exige "
    "que los datos sensibles (entre ellos los de salud) se almacenen con resguardos "
    "tecnicos y organizativos adecuados."
)
parrafo(
    "MedScribe respeta esos lineamientos: las contrasenas se guardan con hash BCrypt, "
    "los datos viajan por HTTPS, el control de acceso se aplica por roles y todas las "
    "acciones criticas quedan registradas en una tabla de auditoria. La estructura de "
    "las notas sigue el formato SOAP, aceptado por la mayor parte de los sistemas "
    "publicos peruanos."
)

salto_pagina()

# =========================
# 4. OBJETIVOS
# =========================
agregar_titulo("4. OBJETIVOS", nivel=1)
parrafo(
    "Los objetivos siguen el criterio SMART: especificos, medibles, alcanzables, "
    "relevantes y con plazo definido."
)

agregar_titulo("Objetivo General", nivel=2)
parrafo(
    "Desarrollar una aplicacion web funcional que reduzca en al menos un 60 % el "
    "tiempo que un medico dedica a escribir notas clinicas, mediante la transcripcion "
    "automatica de la consulta y la generacion estructurada de la nota, durante el "
    "ciclo academico 2026-I del curso DSW1."
)

agregar_titulo("Objetivo Especifico 1", nivel=2)
lista([
    "Especifico: Implementar un pipeline de IA que transcriba audios medicos en menos de 30 segundos para grabaciones de hasta 5 minutos.",
    "Medible: Tiempo de respuesta promedio menor a 30 s.",
    "Alcanzable: Usando Whisper (Groq) y diarizacion en la nube.",
    "Relevante: La transcripcion es el cuello de botella principal del proceso.",
    "Tiempo: Antes del 2 de mayo de 2026.",
])

agregar_titulo("Objetivo Especifico 2", nivel=2)
lista([
    "Especifico: Generar documentos clinicos en formato PDF y Word listos para descarga, con un disenio profesional y secciones SOAP.",
    "Medible: Al menos 2 formatos de salida (PDF y DOCX) por consulta.",
    "Alcanzable: Mediante ReportLab y python-docx con plantillas configurables.",
    "Relevante: Permite que la nota se imprima o se anexe al expediente.",
    "Tiempo: Antes del 2 de mayo de 2026.",
])

agregar_titulo("Objetivo Especifico 3", nivel=2)
lista([
    "Especifico: Implementar control de acceso por roles (Medico y Administrador) con permisos granulares por modulo.",
    "Medible: 6 modulos protegidos con 4 acciones cada uno (ver, crear, editar, eliminar).",
    "Alcanzable: Mediante autenticacion con JWT y permisos almacenados en JSON.",
    "Relevante: Cumple con la rubrica del curso (4 puntos en autenticacion).",
    "Tiempo: Antes del 2 de mayo de 2026.",
])

# =========================
# 5. JUSTIFICACION
# =========================
agregar_titulo("5. JUSTIFICACION DEL PROYECTO", nivel=1)
parrafo(
    "El proyecto contribuye al sector salud de tres formas concretas: (a) devuelve "
    "tiempo de calidad al medico, lo que mejora la relacion humana con el paciente; "
    "(b) reduce errores de transcripcion porque la nota se construye directamente "
    "desde el audio original, sin intermediarios; (c) abarata el costo de la "
    "documentacion clinica al usar modelos gratuitos en la nube."
)

agregar_titulo("Beneficiarios directos", nivel=2)
lista([
    "Medicos generales y especialistas que utilizan la aplicacion para registrar consultas.",
    "Personal administrativo de clinicas que gestiona pacientes y reportes.",
    "Pacientes, que reciben mas atencion verbal y obtienen un documento legible al final de la consulta.",
])

agregar_titulo("Beneficiarios indirectos", nivel=2)
lista([
    "Familiares de los pacientes, que pueden leer la nota clinica entregada en formato PDF.",
    "Instituciones de salud (EsSalud, MINSA, clinicas privadas) que reciben notas estandarizadas en formato SOAP.",
    "Aseguradoras de salud, que procesan reembolsos a partir de notas claras y completas.",
])

salto_pagina()

# =========================
# 6. DEFINICION Y ALCANCE
# =========================
agregar_titulo("6. DEFINICION Y ALCANCE", nivel=1)

agregar_titulo("6.1 Vision general", nivel=2)
parrafo(
    "MedScribe es una aplicacion web multiusuario y multi-clinica. Cada clinica se "
    "registra una sola vez, contrata un plan, agrega a sus medicos y empieza a operar. "
    "La pieza central es el flujo de consulta: el medico abre una nueva consulta, "
    "selecciona al paciente, graba el audio (o lo sube como archivo), revisa la nota "
    "que la inteligencia artificial genera, edita lo que crea conveniente, aprueba y "
    "descarga el PDF o el Word."
)

agregar_titulo("6.2 Arquitectura del sistema", nivel=2)
parrafo(
    "La aplicacion sigue una arquitectura en cuatro capas, separadas y comunicadas por "
    "APIs HTTP. Cada capa puede crecer o cambiar de forma independiente."
)
insertar_imagen("01-arquitectura.png", "Figura 1. Arquitectura general de MedScribe", ancho_cm=15)

parrafo(
    "Capa de presentacion (React + Vite + Tailwind): es lo que el medico ve en su "
    "navegador. Usa componentes reutilizables, hooks personalizados y un contexto "
    "global de autenticacion. Tambien incluye animaciones 3D simples para mejorar la "
    "experiencia visual."
)
parrafo(
    "Capa de negocio (.NET 9 / ASP.NET Core Web API): es la puerta de entrada "
    "(gateway). Recibe las peticiones del frontend, valida los datos, comprueba el "
    "JWT, llama a los stored procedures de SQL Server y orquesta las llamadas al "
    "servicio de IA. Toda la logica de permisos y multi-tenant vive aqui."
)
parrafo(
    "Capa de inteligencia artificial (Python + FastAPI): aqui ocurre la magia. Recibe "
    "audios y transcripciones, ejecuta el pipeline de IA, genera los documentos PDF y "
    "Word, y los devuelve al gateway."
)
parrafo(
    "Capa de datos (SQL Server + Qdrant + Redis + sistema de archivos): SQL Server "
    "guarda usuarios, pacientes, consultas y documentos. Qdrant es una base vectorial "
    "que sirve para buscar contexto medico relevante (RAG). Redis maneja la cola de "
    "trabajos asincronos. El sistema de archivos guarda los PDFs y Words generados."
)

agregar_titulo("6.3 Modulos funcionales", nivel=2)
tabla(
    ["Modulo", "Funcionalidad principal"],
    [
        ["Autenticacion", "Inicio de sesion, registro de clinica, cambio de contrasena."],
        ["Pacientes", "Alta, baja, modificacion y consulta de pacientes por clinica."],
        ["Consultas", "Crear consulta desde audio, revisar nota generada, editar y aprobar."],
        ["Documentos", "Generar PDF y Word, listar, aprobar y descargar."],
        ["Plantillas", "Configurar las secciones que la IA debe rellenar."],
        ["Usuarios", "Gestionar los usuarios de la clinica y asignarles roles."],
        ["Roles", "Crear roles personalizados con permisos granulares por modulo."],
        ["Perfil", "Cada usuario edita sus datos personales y cambia su contrasena."],
    ],
    anchos=[4, 11],
)

agregar_titulo("6.4 Tecnologias utilizadas", nivel=2)
tabla(
    ["Componente", "Tecnologia", "Para que se usa"],
    [
        ["Frontend", "React 19 + TypeScript + Vite + Tailwind CSS + Three.js", "Interfaz del usuario en el navegador."],
        ["Backend", "C# .NET 9 + ASP.NET Core Web API", "Gateway REST, logica de negocio y seguridad."],
        ["Servicio IA", "Python 3.12 + FastAPI", "Pipeline de transcripcion y generacion de documentos."],
        ["Base relacional", "SQL Server Express", "Datos transaccionales (usuarios, consultas, etc.)."],
        ["Base vectorial", "Qdrant", "Busqueda de contexto medico (RAG)."],
        ["Cola y cache", "Redis", "Trabajos asincronos e idempotencia."],
        ["Acceso a datos", "ADO.NET + Stored Procedures", "Acceso seguro y rapido a SQL Server."],
        ["Autenticacion", "JWT + BCrypt", "Sesiones seguras y contrasenas hasheadas."],
        ["Transcripcion", "Whisper (Groq API)", "Convertir voz en texto."],
        ["Diarizacion", "Pyannote 3.1 + Deepgram Nova-3", "Identificar quien habla en el audio."],
        ["LLM", "Llama 3.3 70B (Groq API)", "Generar la nota clinica estructurada."],
        ["Documentos", "ReportLab + python-docx", "Generar archivos PDF y Word."],
        ["Empaquetado", "Docker + Docker Compose", "Levantar todo el sistema con un comando."],
    ],
    anchos=[3, 6, 6],
)

agregar_titulo("6.5 Diagramas RUP", nivel=2)
parrafo(
    "A continuacion se presentan los diagramas que documentan el sistema, siguiendo "
    "los lineamientos de la metodologia RUP (Rational Unified Process)."
)

agregar_titulo("6.5.1 Diagrama de Casos de Uso General", nivel=3)
parrafo(
    "El sistema reconoce dos actores: el Medico, que utiliza la aplicacion para "
    "atender a sus pacientes, y el Administrador, que ademas configura la clinica, "
    "los usuarios y los roles."
)
insertar_imagen("02-casos-uso.png", "Figura 2. Diagrama de Casos de Uso general", ancho_cm=15)

agregar_titulo("6.5.2 Diagrama de Clases", nivel=3)
parrafo(
    "Muestra las clases principales del dominio y como se relacionan. Una Clinica "
    "agrupa Usuarios y Pacientes; un Usuario puede ser Medico; un Medico atiende "
    "Consultas; cada Consulta genera Documentos y rellena Valores de Seccion "
    "siguiendo una Plantilla."
)
insertar_imagen("03-clases.png", "Figura 3. Diagrama de Clases del dominio", ancho_cm=15)

agregar_titulo("6.5.3 Modelo Entidad-Relacion", nivel=3)
parrafo(
    "Vista fisica de la base de datos en SQL Server. Cada tabla incluye su clave "
    "primaria, sus campos principales y las relaciones con las demas tablas."
)
insertar_imagen("04-er-base-datos.png", "Figura 4. Modelo Entidad-Relacion", ancho_cm=15)

agregar_titulo("6.5.4 Diagramas de Secuencia", nivel=3)
parrafo(
    "Cada diagrama de secuencia describe paso a paso como se comunican los "
    "componentes para realizar una operacion concreta. Mostramos los dos casos de "
    "uso mas importantes: iniciar sesion y crear una consulta a partir de audio."
)
insertar_imagen("05-secuencia-login.png", "Figura 5. Secuencia de inicio de sesion", ancho_cm=15)
insertar_imagen("06-secuencia-transcripcion.png", "Figura 6. Secuencia de creacion de consulta con IA", ancho_cm=15)

agregar_titulo("6.5.5 Diagrama de Actividades", nivel=3)
parrafo(
    "Visualiza el flujo completo del pipeline de inteligencia artificial: desde que "
    "el medico graba el audio hasta que el documento queda aprobado y archivado."
)
insertar_imagen("07-actividades-pipeline.png", "Figura 7. Diagrama de actividades del pipeline IA", ancho_cm=15)

agregar_titulo("6.5.6 Diagrama de Componentes", nivel=3)
parrafo(
    "Muestra como esta organizado el codigo de cada capa y como se comunican entre "
    "si las distintas piezas del sistema."
)
insertar_imagen("08-componentes.png", "Figura 8. Diagrama de Componentes", ancho_cm=15)

agregar_titulo("6.5.7 Diagrama de Despliegue", nivel=3)
parrafo(
    "Indica como se distribuyen los servicios en contenedores Docker, sus puertos y "
    "las APIs externas con las que el sistema se comunica."
)
insertar_imagen("09-despliegue.png", "Figura 9. Diagrama de Despliegue", ancho_cm=15)

agregar_titulo("6.5.8 Diagrama de Estados de la Consulta", nivel=3)
parrafo(
    "Toda Consulta pasa por varios estados a lo largo de su vida util. Este diagrama "
    "indica las transiciones permitidas y los eventos que las disparan."
)
insertar_imagen("10-estados-consulta.png", "Figura 10. Diagrama de Estados de la Consulta", ancho_cm=14)

salto_pagina()

# =========================
# 7. PRODUCTOS Y ENTREGABLES
# =========================
agregar_titulo("7. PRODUCTOS Y ENTREGABLES", nivel=1)
parrafo(
    "Como cierre del proyecto, el equipo entrega los siguientes artefactos:"
)
tabla(
    ["Entregable", "Formato", "Ubicacion"],
    [
        ["Codigo fuente completo del sistema", "Repositorio Git + ZIP", "github.com/roberto3101/MEDSCRIBE-DSW1-PILOTO"],
        ["Informe del proyecto", "Word (.docx)", "Documentos/Negocio/"],
        ["Especificaciones de Casos de Uso", "Word (.docx)", "Documentos/Tecnico/"],
        ["Diagramas RUP", "PNG + SVG + Mermaid (.mmd)", "Documentos/Tecnico/Diagramas/"],
        ["Presentacion interactiva (web)", "HTML responsivo offline", "Documentos/Web/"],
        ["Documento de estructura del proyecto", "Markdown (.md)", "Documentos/Tecnico/"],
        ["Manual de instalacion y uso", "README.md", "Raiz del repositorio"],
        ["Coleccion Postman para pruebas", "JSON", "postman/"],
    ],
    anchos=[5.5, 3, 6.5],
)

# =========================
# 8. CONCLUSIONES
# =========================
agregar_titulo("8. CONCLUSIONES", nivel=1)
parrafo(
    "1. La combinacion de modelos de inteligencia artificial gratuitos (Whisper para "
    "transcripcion, Llama para estructuracion) con un backend tradicional en .NET y "
    "una base SQL Server permite construir productos clinicos viables sin costos de "
    "licenciamiento prohibitivos."
)
parrafo(
    "2. La separacion del sistema en cuatro capas independientes (frontend, gateway, "
    "servicio IA y base de datos) facilito el trabajo en paralelo del equipo y "
    "permitio que cada integrante se especializara en una tecnologia distinta sin "
    "bloquear a los demas."
)
parrafo(
    "3. El uso de Docker Compose elimina por completo el problema clasico de las "
    "demos universitarias (\"en mi computadora si funciona\") y deja el sistema "
    "listo para evaluar en cualquier maquina con un solo comando."
)

# =========================
# 9. RECOMENDACIONES
# =========================
agregar_titulo("9. RECOMENDACIONES", nivel=1)
parrafo(
    "1. Para equipos que quieran continuar este proyecto se recomienda agregar "
    "soporte multi-idioma en la interfaz, ya que el sistema actual esta en espanol "
    "pero la IA puede transcribir en cualquier idioma."
)
parrafo(
    "2. Conviene agregar pruebas automatizadas (unit tests + e2e con Playwright) "
    "antes de incorporar nuevas funcionalidades, para evitar regresiones cuando se "
    "modifique el pipeline de IA."
)
parrafo(
    "3. Si el sistema se quiere usar en produccion real, se debe implementar "
    "anonimizacion de datos antes de enviarlos a la nube y obtener la certificacion "
    "del MINSA segun el reglamento RENHICE."
)

# =========================
# 10. GLOSARIO
# =========================
agregar_titulo("10. GLOSARIO", nivel=1)
tabla(
    ["Termino", "Definicion"],
    [
        ["API REST", "Forma estandar de comunicacion entre programas usando HTTP y JSON."],
        ["BCrypt", "Algoritmo para guardar contrasenas de forma segura, dificil de revertir."],
        ["Diarizacion", "Identificar automaticamente quien habla en una grabacion de audio."],
        ["Docker", "Tecnologia que empaqueta un programa y todo lo que necesita en un contenedor."],
        ["FastAPI", "Framework moderno de Python para construir APIs rapidas."],
        ["JWT", "Token firmado que demuestra que el usuario esta autenticado."],
        ["LLM", "Large Language Model. Modelo de inteligencia artificial entrenado para entender y generar texto."],
        ["Multi-tenant", "Aplicacion en la que varias organizaciones (clinicas) usan la misma instalacion sin verse entre si."],
        ["Pipeline", "Cadena de pasos en orden donde la salida de uno es la entrada del siguiente."],
        ["Pyannote", "Libreria de Python para reconocimiento e identificacion de voces."],
        ["RAG", "Retrieval Augmented Generation. Tecnica que da a un LLM contexto adicional buscado en una base externa."],
        ["RBAC", "Role Based Access Control. Sistema de permisos basado en roles."],
        ["RUP", "Rational Unified Process. Metodologia formal de desarrollo de software."],
        ["SOAP (medico)", "Formato de nota clinica con secciones Subjetivo, Objetivo, Analisis y Plan."],
        ["Stored Procedure", "Programa almacenado dentro de la base de datos que ejecuta una operacion completa."],
        ["Whisper", "Modelo de OpenAI que convierte audio en texto."],
    ],
    anchos=[4, 11],
)

# =========================
# 11. BIBLIOGRAFIA
# =========================
agregar_titulo("11. BIBLIOGRAFIA", nivel=1)
referencias = [
    "Colegio Medico del Peru. (2024). Estadisticas del personal medico colegiado. Recuperado de https://www.cmp.org.pe",
    "Instituto Nacional de Estadistica e Informatica. (2023). Encuesta Nacional de Satisfaccion de Usuarios en Salud. INEI.",
    "Ministerio de Salud del Peru. (2017). Decreto Supremo N.° 009-2017-SA. Reglamento de la Ley N.° 30024. Diario El Peruano.",
    "Ministerio de Salud del Peru. (2023). Diagnostico de la digitalizacion de consultorios privados en Peru. MINSA.",
    "OpenAI. (2022). Robust Speech Recognition via Large Scale Weak Supervision (Whisper). https://github.com/openai/whisper",
    "Microsoft. (2024). Documentacion oficial de ASP.NET Core 9. https://learn.microsoft.com/aspnet/core",
    "Microsoft. (2024). Documentacion oficial de SQL Server 2022. https://learn.microsoft.com/sql",
    "Tiangolo, S. (2024). FastAPI Documentation. https://fastapi.tiangolo.com",
    "Meta AI. (2024). Llama 3.3 Model Card. https://llama.meta.com",
    "Bredin, H. et al. (2023). Pyannote.audio: Neural building blocks for speaker diarization. https://github.com/pyannote/pyannote-audio",
    "Qdrant. (2024). Vector Database Documentation. https://qdrant.tech/documentation",
    "Kruchten, P. (2003). The Rational Unified Process: An Introduction (3rd ed.). Addison-Wesley.",
]
for ref in referencias:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(ref)
    run.font.name = "Arial"
    run.font.size = Pt(10)

# =========================
# 12. ANEXOS
# =========================
salto_pagina()
agregar_titulo("12. ANEXOS", nivel=1)

agregar_titulo("Anexo A. Estructura de carpetas del proyecto", nivel=2)
parrafo(
    "El repositorio esta organizado de la siguiente forma:", justificar=False
)
estructura = (
    "MEDSCRIBE-DSW1-PILOTO/\n"
    "  cliente-web/         (Frontend React + Vite + Tailwind)\n"
    "  gateway-dotnet/      (Backend C# .NET 9 - Gateway API)\n"
    "  servicio-ia/         (Servicio Python FastAPI con IA)\n"
    "  base-datos/          (Migraciones SQL Server y stored procedures)\n"
    "  postman/             (Colecciones de prueba HTTP)\n"
    "  Documentos/          (Toda la documentacion del proyecto)\n"
    "    Negocio/           (Informe principal en Word)\n"
    "    Tecnico/           (ECUs, diagramas y estructura)\n"
    "      Diagramas/       (PNG + SVG + fuentes Mermaid)\n"
    "    Web/               (Presentacion interactiva HTML offline)\n"
    "    _scripts/          (Scripts Python para regenerar todo)\n"
    "  docker-compose.yml   (Orquestacion de contenedores)\n"
    "  README.md            (Manual de instalacion rapida)\n"
)
p = doc.add_paragraph(estructura)
p.runs[0].font.name = "Consolas"
p.runs[0].font.size = Pt(9)

agregar_titulo("Anexo B. Distribucion del trabajo del equipo", nivel=2)
tabla(
    ["Integrante", "Codigo", "Responsabilidad principal", "Que expone"],
    [
        ["Jose Roberto La Rosa Ledezma", "i202333980", "Coordinacion general + Servicio IA Python", "Servicio IA Python (FastAPI, Whisper, Llama, RAG)"],
        ["Jason Davila Delgado", "i202415540", "Backend Gateway en C# .NET 9", "Backend C# (.NET, JWT, ADO.NET, controladores)"],
        ["Luis Joel Curi", "i202417794", "Base de datos SQL Server", "Modelo ER, stored procedures, transacciones"],
        ["Edward Alexander Escobedo Murga", "i201917851", "Frontend React + Arquitectura general", "Frontend React, Tailwind, Docker, integracion"],
    ],
    anchos=[4.5, 2, 4.5, 5],
)

agregar_titulo("Anexo C. Comando de instalacion rapida", nivel=2)
parrafo(
    "Cualquier persona puede levantar el sistema completo con tres comandos:",
    justificar=False,
)
comandos = (
    "git clone https://github.com/roberto3101/MEDSCRIBE-DSW1-PILOTO.git\n"
    "cd MEDSCRIBE-DSW1-PILOTO\n"
    "docker compose up --build\n"
)
p = doc.add_paragraph(comandos)
p.runs[0].font.name = "Consolas"
p.runs[0].font.size = Pt(10)

parrafo(
    "Una vez que los contenedores se levanten, la aplicacion estara disponible en "
    "http://localhost:3000."
)

doc.save(SALIDA)
print(f"OK: {SALIDA}")
print(f"Tamano: {SALIDA.stat().st_size // 1024} KB")
