from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage
import io
import os
import base64
import tempfile
from datetime import datetime


_FORMATOS_NATIVOS_DOCX = {"png", "jpg", "jpeg", "bmp", "gif", "tiff", "tif"}


def _imagen_compatible_con_docx(ruta: str):
    """python-docx solo acepta PNG/JPG/BMP/GIF/TIFF. Si el archivo es WebP/SVG/HEIC,
    lo convertimos a PNG en memoria y devolvemos un BufferedReader compatible."""
    if not ruta or not os.path.exists(ruta):
        return ruta
    extension = ruta.rsplit(".", 1)[-1].lower() if "." in ruta else ""
    if extension in _FORMATOS_NATIVOS_DOCX:
        return ruta
    try:
        with PILImage.open(ruta) as img:
            modo_destino = "RGBA" if img.mode in ("RGBA", "LA", "P") else "RGB"
            buffer = io.BytesIO()
            img.convert(modo_destino).save(buffer, format="PNG")
            buffer.seek(0)
            return buffer
    except Exception:
        return ruta


NAVY = RGBColor(0x0B, 0x2B, 0x4C)
NAVY_HEX = "0B2B4C"
ACENTO_DORADO = RGBColor(0xB0, 0x8D, 0x57)
ACENTO_DORADO_HEX = "B08D57"
GRIS_BORDE_HEX = "D4D4D8"
GRIS_BORDE_SUAVE_HEX = "E4E4E7"
TEXTO_PRIMARIO = RGBColor(0x18, 0x20, 0x2C)
TEXTO_SUAVE = RGBColor(0x57, 0x63, 0x71)
TEXTO_META = RGBColor(0x71, 0x7A, 0x85)
FONDO_META_HEX = "FAFAF7"


_archivos_temporales: list[str] = []


def _base64_a_ruta_temporal(data_url: str) -> str:
    if not data_url or not data_url.startswith("data:image"):
        return ""
    try:
        _, datos = data_url.split(",", 1)
        ruta = tempfile.NamedTemporaryFile(suffix=".png", delete=False).name
        with open(ruta, "wb") as f:
            f.write(base64.b64decode(datos))
        _archivos_temporales.append(ruta)
        return ruta
    except Exception:
        return ""


def _limpiar_temporales() -> None:
    global _archivos_temporales
    for ruta in _archivos_temporales:
        if ruta and os.path.exists(ruta):
            try:
                os.unlink(ruta)
            except Exception:
                pass
    _archivos_temporales = []


def _sombreado_celda(celda, color_hex: str) -> None:
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def _bordes_tabla(tabla, color_hex: str = GRIS_BORDE_HEX, tamano: str = "4",
                  bordes: tuple = ("top", "left", "bottom", "right", "insideH", "insideV")) -> None:
    tbl = tabla._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for tipo in bordes:
        borde = OxmlElement(f'w:{tipo}')
        borde.set(qn('w:val'), 'single')
        borde.set(qn('w:sz'), tamano)
        borde.set(qn('w:color'), color_hex)
        tblBorders.append(borde)
    tblPr.append(tblBorders)


def _solo_borde(tabla, lado: str, color_hex: str, tamano: str = "4") -> None:
    tbl = tabla._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for tipo in ("top", "left", "bottom", "right"):
        borde = OxmlElement(f'w:{tipo}')
        if tipo == lado:
            borde.set(qn('w:val'), 'single')
            borde.set(qn('w:sz'), tamano)
            borde.set(qn('w:color'), color_hex)
        else:
            borde.set(qn('w:val'), 'nil')
        tblBorders.append(borde)
    tblPr.append(tblBorders)


def _aplicar_small_caps_y_espaciado(run, espaciado_puntos: float = 2.0) -> None:
    rPr = run._element.get_or_add_rPr()
    caps = OxmlElement('w:smallCaps')
    caps.set(qn('w:val'), '1')
    rPr.append(caps)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), str(int(espaciado_puntos * 20)))
    rPr.append(spacing)


def _aplicar_letterspacing(run, espaciado_puntos: float) -> None:
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), str(int(espaciado_puntos * 20)))
    rPr.append(spacing)


def _regla_horizontal(doc: Document, grosor_puntos: float = 0.5, color_hex: str = GRIS_BORDE_HEX,
                      espacio_antes: int = 0, espacio_despues: int = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(espacio_antes)
    p.paragraph_format.space_after = Pt(espacio_despues)
    pPr = p.paragraph_format.element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(grosor_puntos * 8)))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _regla_doble(doc: Document, color_hex: str = NAVY_HEX, espacio_despues: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(espacio_despues)
    pPr = p.paragraph_format.element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'double')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _configurar_documento(doc: Document) -> None:
    for seccion in doc.sections:
        seccion.top_margin = Cm(2.2)
        seccion.bottom_margin = Cm(1.8)
        seccion.left_margin = Cm(2.3)
        seccion.right_margin = Cm(2.3)

    estilo_normal = doc.styles['Normal']
    estilo_normal.font.name = 'Cambria'
    rPr = estilo_normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Cambria')
    rFonts.set(qn('w:hAnsi'), 'Cambria')
    rFonts.set(qn('w:cs'), 'Cambria')
    estilo_normal.font.size = Pt(10.5)
    estilo_normal.font.color.rgb = TEXTO_PRIMARIO
    estilo_normal.paragraph_format.line_spacing = 1.25
    estilo_normal.paragraph_format.space_after = Pt(3)


def _establecer_fuente(run, fuente: str) -> None:
    run.font.name = fuente
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), fuente)
    rFonts.set(qn('w:hAnsi'), fuente)
    rFonts.set(qn('w:cs'), fuente)


def _agregar_encabezado(doc: Document, config: dict) -> None:
    nombre_clinica = config.get("nombre_clinica", "MedScribe AI")
    ruc = config.get("ruc", "")
    direccion = config.get("direccion", "")
    telefono = config.get("telefono", "")
    correo = config.get("correo", "")
    logo_path = config.get("logo_path", "")
    tiene_logo = logo_path and os.path.exists(logo_path)

    fecha = datetime.now().strftime("%d/%m/%Y")
    hora = datetime.now().strftime("%H:%M")
    numero = config.get("numero_documento", f"HC-{datetime.now().strftime('%Y%m%d%H%M%S')}")

    if tiene_logo:
        tabla = doc.add_table(rows=1, cols=3)
        anchos = [Cm(2.2), Cm(10.5), Cm(4.3)]
    else:
        tabla = doc.add_table(rows=1, cols=2)
        anchos = [Cm(12.7), Cm(4.3)]

    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.autofit = False
    for i, ancho in enumerate(anchos):
        tabla.columns[i].width = ancho
        for celda in tabla.columns[i].cells:
            celda.width = ancho
    for fila in tabla.rows:
        for celda in fila.cells:
            celda.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    indice = 0
    if tiene_logo:
        celda_logo = tabla.cell(0, 0)
        p_logo = celda_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after = Pt(0)
        p_logo.paragraph_format.line_spacing = 1.0
        try:
            p_logo.add_run().add_picture(_imagen_compatible_con_docx(logo_path), width=Cm(1.8))
        except Exception:
            pass
        indice = 1

    celda_info = tabla.cell(0, indice)
    p_nombre = celda_info.paragraphs[0]
    p_nombre.paragraph_format.space_after = Pt(2)
    run = p_nombre.add_run(nombre_clinica.upper())
    _establecer_fuente(run, 'Cambria')
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    run.font.bold = True
    _aplicar_letterspacing(run, 1.5)

    p_lema = celda_info.add_paragraph()
    p_lema.paragraph_format.space_after = Pt(4)
    run_lema = p_lema.add_run("Historia clinica electronica institucional")
    _establecer_fuente(run_lema, 'Cambria')
    run_lema.font.size = Pt(8)
    run_lema.font.italic = True
    run_lema.font.color.rgb = TEXTO_SUAVE

    lineas_info = []
    if ruc:
        lineas_info.append(f"RUC {ruc}")
    if direccion:
        lineas_info.append(direccion)
    partes_contacto = []
    if telefono:
        partes_contacto.append(f"Tel. {telefono}")
    if correo:
        partes_contacto.append(correo)
    if partes_contacto:
        lineas_info.append("  ·  ".join(partes_contacto))

    if lineas_info:
        p = celda_info.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run_info = p.add_run("   ·   ".join(lineas_info))
        _establecer_fuente(run_info, 'Calibri')
        run_info.font.size = Pt(8)
        run_info.font.color.rgb = TEXTO_META

    celda_meta = tabla.cell(0, indice + 1)
    _sombreado_celda(celda_meta, FONDO_META_HEX)
    primer_parrafo_meta = True
    for etiqueta, valor in [("FECHA", fecha), ("HORA", hora), ("FOLIO", numero)]:
        if primer_parrafo_meta:
            p = celda_meta.paragraphs[0]
            primer_parrafo_meta = False
        else:
            p = celda_meta.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        run_lbl = p.add_run(f"{etiqueta}  ")
        _establecer_fuente(run_lbl, 'Calibri')
        run_lbl.font.size = Pt(7)
        run_lbl.font.bold = True
        run_lbl.font.color.rgb = ACENTO_DORADO
        _aplicar_letterspacing(run_lbl, 0.8)
        run_val = p.add_run(valor)
        _establecer_fuente(run_val, 'Calibri')
        run_val.font.size = Pt(9)
        run_val.font.color.rgb = TEXTO_PRIMARIO
        run_val.font.bold = True

    _bordes_tabla(tabla, GRIS_BORDE_HEX, "4", bordes=("top", "left", "bottom", "right"))

    _regla_doble(doc, NAVY_HEX, espacio_despues=2)
    regla_acento = doc.add_paragraph()
    regla_acento.paragraph_format.space_before = Pt(0)
    regla_acento.paragraph_format.space_after = Pt(8)
    pPr = regla_acento.paragraph_format.element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), ACENTO_DORADO_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _agregar_titulo(doc: Document, tipo_documento: str) -> None:
    titulos = {
        "SOAP": "Nota Clinica SOAP",
        "HistoriaClinica": "Historia Clinica",
        "Receta": "Receta Medica",
        "Personalizada": "Documento Clinico",
    }
    titulo = titulos.get(tipo_documento, "Documento Clinico")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(titulo.upper())
    _establecer_fuente(run, 'Cambria')
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    run.font.bold = True
    _aplicar_letterspacing(run, 4)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(4)
    run_sub = p_sub.add_run("Documento medico oficial  ·  Firmado digitalmente")
    _establecer_fuente(run_sub, 'Cambria')
    run_sub.font.size = Pt(8)
    run_sub.font.italic = True
    run_sub.font.color.rgb = TEXTO_SUAVE
    _aplicar_letterspacing(run_sub, 0.6)

    linea = doc.add_paragraph()
    linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    linea.paragraph_format.space_before = Pt(0)
    linea.paragraph_format.space_after = Pt(14)
    run_linea = linea.add_run("◆  ◆  ◆")
    _establecer_fuente(run_linea, 'Cambria')
    run_linea.font.size = Pt(7)
    run_linea.font.color.rgb = ACENTO_DORADO


def _agregar_titulo_seccion(doc: Document, titulo: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True

    run_barra = p.add_run("│ ")
    _establecer_fuente(run_barra, 'Cambria')
    run_barra.font.size = Pt(12)
    run_barra.font.color.rgb = ACENTO_DORADO
    run_barra.font.bold = True

    run = p.add_run(titulo.upper())
    _establecer_fuente(run, 'Cambria')
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.font.bold = True
    _aplicar_letterspacing(run, 2)

    regla = doc.add_paragraph()
    regla.paragraph_format.space_before = Pt(1)
    regla.paragraph_format.space_after = Pt(4)
    regla.paragraph_format.keep_with_next = True
    pPr = regla.paragraph_format.element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), GRIS_BORDE_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _agregar_contenido_sin_campos_vacios(doc: Document, nota_clinica: str) -> None:
    seccion_actual = None
    contenido_seccion = []

    def _flush_seccion():
        nonlocal seccion_actual, contenido_seccion
        if seccion_actual is None:
            return

        tiene_contenido_real = any(
            linea.strip() and
            "no se proporcion" not in linea.lower() and
            "no se tiene" not in linea.lower() and
            "no se registr" not in linea.lower() and
            "no se mencion" not in linea.lower() and
            "no disponible" not in linea.lower()
            for linea in contenido_seccion
        )

        if not tiene_contenido_real:
            seccion_actual = None
            contenido_seccion = []
            return

        _agregar_titulo_seccion(doc, seccion_actual)

        for linea in contenido_seccion:
            if not linea.strip():
                continue
            if linea.startswith("- ") or linea.startswith("• "):
                texto_bullet = linea[2:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.2
                run_bullet = p.add_run("▸  ")
                _establecer_fuente(run_bullet, 'Cambria')
                run_bullet.font.size = Pt(10)
                run_bullet.font.color.rgb = ACENTO_DORADO
                run_bullet.font.bold = True
                run_texto = p.add_run(texto_bullet)
                _establecer_fuente(run_texto, 'Cambria')
                run_texto.font.size = Pt(10.5)
                run_texto.font.color.rgb = TEXTO_PRIMARIO
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.3
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(linea)
                _establecer_fuente(run, 'Cambria')
                run.font.size = Pt(10.5)
                run.font.color.rgb = TEXTO_PRIMARIO

        espaciador = doc.add_paragraph()
        espaciador.paragraph_format.space_after = Pt(2)

        seccion_actual = None
        contenido_seccion = []

    for linea in nota_clinica.split("\n"):
        texto = linea.strip()

        es_titulo = False
        titulo_limpio = ""

        if texto.startswith("## "):
            titulo_limpio = texto[3:].strip().strip("*")
            es_titulo = True
        elif texto.startswith("# "):
            titulo_limpio = texto[2:].strip().strip("*")
            es_titulo = True
        elif texto.startswith("**") and texto.endswith("**"):
            titulo_limpio = texto.strip("*").strip()
            es_titulo = True

        if es_titulo:
            _flush_seccion()
            seccion_actual = titulo_limpio
            contenido_seccion = []
        elif seccion_actual is not None:
            contenido_seccion.append(texto)
        elif texto:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.3
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(texto)
            _establecer_fuente(run, 'Cambria')
            run.font.size = Pt(10.5)
            run.font.color.rgb = TEXTO_PRIMARIO

    _flush_seccion()


def _establecer_altura_fila(fila, alto_cm: float, regla: str = "exact") -> None:
    trPr = fila._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(alto_cm * 567)))
    trHeight.set(qn('w:hRule'), regla)
    trPr.append(trHeight)


def _rellenar_celda_firma_grafica(celda, firma_path: str) -> None:
    celda.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    primero = celda.paragraphs[0]
    primero.alignment = WD_ALIGN_PARAGRAPH.CENTER
    primero.paragraph_format.space_before = Pt(0)
    primero.paragraph_format.space_after = Pt(0)
    primero.paragraph_format.line_spacing = 1.0

    if firma_path:
        try:
            primero.add_run().add_picture(_imagen_compatible_con_docx(firma_path), width=Cm(4.8), height=Cm(1.6))
            return
        except Exception:
            pass
    run = primero.add_run("")


def _rellenar_celda_texto_firma(celda, rol: str, nombre: str, detalles: list[str]) -> None:
    celda.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tcPr = celda._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:color'), "3F4A55")
    tcBorders.append(top)
    for lado in ("left", "bottom", "right"):
        b = OxmlElement(f'w:{lado}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    primero = celda.paragraphs[0]
    primero.alignment = WD_ALIGN_PARAGRAPH.CENTER
    primero.paragraph_format.space_before = Pt(3)
    primero.paragraph_format.space_after = Pt(1)
    if rol:
        run_rol = primero.add_run(rol.upper())
        _establecer_fuente(run_rol, 'Cambria')
        run_rol.font.size = Pt(7.5)
        run_rol.font.color.rgb = ACENTO_DORADO
        _aplicar_letterspacing(run_rol, 1.2)

    if nombre:
        p_nombre = celda.add_paragraph()
        p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_nombre.paragraph_format.space_after = Pt(1)
        run = p_nombre.add_run(nombre)
        _establecer_fuente(run, 'Cambria')
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY

    for detalle in detalles:
        if not detalle:
            continue
        p = celda.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(detalle)
        _establecer_fuente(run, 'Cambria')
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = TEXTO_SUAVE


def _agregar_firmas(doc: Document, config: dict) -> None:
    firma_medico_ruta = _base64_a_ruta_temporal(config.get("firma_medico", ""))
    firma_clinica_ruta = _base64_a_ruta_temporal(config.get("firma_clinica", ""))

    nombre_medico = config.get("nombre_medico", "")
    colegiatura = config.get("colegiatura", "")
    especialidad = config.get("especialidad_medico", "")
    nombre_clinica = config.get("nombre_clinica", "")

    if not nombre_medico and not nombre_clinica and not firma_medico_ruta and not firma_clinica_ruta:
        return

    espaciador = doc.add_paragraph()
    espaciador.paragraph_format.space_before = Pt(18)
    espaciador.paragraph_format.space_after = Pt(4)

    tabla = doc.add_table(rows=2, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.autofit = False
    ancho_columna = Cm(8.3)
    for i in range(2):
        tabla.columns[i].width = ancho_columna
        for celda in tabla.columns[i].cells:
            celda.width = ancho_columna

    _establecer_altura_fila(tabla.rows[0], 2.0, regla="exact")

    _rellenar_celda_firma_grafica(tabla.cell(0, 0), firma_medico_ruta)
    _rellenar_celda_firma_grafica(tabla.cell(0, 1), firma_clinica_ruta)

    detalles_medico = []
    if colegiatura:
        detalles_medico.append(f"Colegiatura Medica  ·  {colegiatura}")
    if especialidad:
        detalles_medico.append(especialidad)
    _rellenar_celda_texto_firma(
        tabla.cell(1, 0),
        "Medico Tratante",
        nombre_medico,
        detalles_medico,
    )

    detalles_clinica = []
    if config.get("ruc"):
        detalles_clinica.append(f"RUC {config['ruc']}")
    detalles_clinica.append("Sello institucional")
    _rellenar_celda_texto_firma(
        tabla.cell(1, 1),
        "Institucion prestadora",
        nombre_clinica,
        detalles_clinica,
    )

    for fila in tabla.rows:
        for celda in fila.cells:
            if celda is tabla.cell(0, 0) or celda is tabla.cell(0, 1):
                tcPr = celda._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for lado in ("top", "left", "bottom", "right"):
                    b = OxmlElement(f'w:{lado}')
                    b.set(qn('w:val'), 'nil')
                    tcBorders.append(b)
                tcPr.append(tcBorders)


def _agregar_pie(doc: Document) -> None:
    espaciador = doc.add_paragraph()
    espaciador.paragraph_format.space_before = Pt(14)
    espaciador.paragraph_format.space_after = Pt(0)

    _regla_horizontal(doc, grosor_puntos=0.5, color_hex=GRIS_BORDE_HEX, espacio_despues=4)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(1)
    run1 = p1.add_run("DOCUMENTO MEDICO OFICIAL")
    _establecer_fuente(run1, 'Cambria')
    run1.font.size = Pt(7.5)
    run1.font.bold = True
    run1.font.color.rgb = NAVY
    _aplicar_letterspacing(run1, 2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(1)
    run2 = p2.add_run(
        f"Generado por MedScribe AI  ·  {datetime.now().strftime('%d/%m/%Y a las %H:%M')}  ·  "
        "Requiere revision y aprobacion del medico tratante"
    )
    _establecer_fuente(run2, 'Cambria')
    run2.font.size = Pt(7.5)
    run2.font.color.rgb = TEXTO_SUAVE
    run2.font.italic = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(0)
    run3 = p3.add_run("Conforme a NTS N.o 139-MINSA/2018/DGAIN  ·  Retencion minima de 20 anios")
    _establecer_fuente(run3, 'Cambria')
    run3.font.size = Pt(7)
    run3.font.color.rgb = TEXTO_META


def _calcular_edad(fecha_nacimiento_iso: str) -> str:
    if not fecha_nacimiento_iso:
        return ""
    try:
        fecha = fecha_nacimiento_iso.split("T")[0]
        partes = fecha.split("-")
        if len(partes) < 3:
            return ""
        anio, mes, dia = int(partes[0]), int(partes[1]), int(partes[2])
        hoy = datetime.now()
        edad = hoy.year - anio - (1 if (hoy.month, hoy.day) < (mes, dia) else 0)
        return f"{edad} anios" if edad >= 0 else ""
    except Exception:
        return ""


def _formatear_fecha_iso(fecha_iso: str) -> str:
    if not fecha_iso:
        return ""
    try:
        fecha = fecha_iso.split("T")[0]
        partes = fecha.split("-")
        if len(partes) == 3:
            return f"{partes[2]}/{partes[1]}/{partes[0]}"
        return fecha
    except Exception:
        return fecha_iso


def _agregar_datos_paciente(doc: Document, config: dict) -> None:
    paciente = config.get("paciente") or {}
    if not paciente.get("nombre_completo") and not paciente.get("numero_documento"):
        return

    nombre = paciente.get("nombre_completo", "").strip()
    tipo_doc = paciente.get("tipo_documento", "DNI").strip() or "DNI"
    num_doc = paciente.get("numero_documento", "").strip()
    sexo = paciente.get("sexo", "").strip()
    fecha_nac = _formatear_fecha_iso(paciente.get("fecha_nacimiento", ""))
    edad = _calcular_edad(paciente.get("fecha_nacimiento", ""))
    telefono = paciente.get("telefono", "").strip()
    correo = paciente.get("correo", "").strip()
    direccion = paciente.get("direccion", "").strip()
    especialidad = config.get("especialidad_consulta", "").strip()

    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(4)
    p_titulo.paragraph_format.space_after = Pt(0)
    p_titulo.paragraph_format.keep_with_next = True
    run_barra = p_titulo.add_run("│ ")
    _establecer_fuente(run_barra, 'Cambria')
    run_barra.font.size = Pt(12)
    run_barra.font.color.rgb = ACENTO_DORADO
    run_barra.font.bold = True
    run_t = p_titulo.add_run("DATOS DEL PACIENTE")
    _establecer_fuente(run_t, 'Cambria')
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = NAVY
    run_t.font.bold = True
    _aplicar_letterspacing(run_t, 2)

    regla = doc.add_paragraph()
    regla.paragraph_format.space_before = Pt(1)
    regla.paragraph_format.space_after = Pt(4)
    regla.paragraph_format.keep_with_next = True
    pPr = regla.paragraph_format.element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), GRIS_BORDE_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)

    tabla = doc.add_table(rows=1, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.autofit = False
    ancho_col = Cm(8.3)
    for i in range(2):
        tabla.columns[i].width = ancho_col
        for celda in tabla.columns[i].cells:
            celda.width = ancho_col
    for celda in tabla.rows[0].cells:
        celda.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _sombreado_celda(celda, FONDO_META_HEX)

    def _insertar_campo(celda, etiqueta: str, valor: str, primero: bool) -> None:
        if primero:
            p = celda.paragraphs[0]
        else:
            p = celda.add_paragraph()
        p.paragraph_format.space_before = Pt(2) if primero else Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run_lbl = p.add_run(f"{etiqueta}  ")
        _establecer_fuente(run_lbl, 'Calibri')
        run_lbl.font.size = Pt(7.5)
        run_lbl.font.bold = True
        run_lbl.font.color.rgb = ACENTO_DORADO
        _aplicar_letterspacing(run_lbl, 0.8)
        run_val = p.add_run(valor or "—")
        _establecer_fuente(run_val, 'Cambria')
        run_val.font.size = Pt(10)
        run_val.font.color.rgb = TEXTO_PRIMARIO

    campos_izquierda = [
        ("NOMBRE", nombre or "—"),
        ("DOCUMENTO", f"{tipo_doc} {num_doc}".strip() if num_doc else "—"),
        ("FECHA DE NACIMIENTO", f"{fecha_nac}  ·  {edad}".strip("  ·  ") if fecha_nac else edad or "—"),
        ("SEXO", sexo or "—"),
    ]

    campos_derecha = []
    if especialidad:
        campos_derecha.append(("ESPECIALIDAD", especialidad))
    if telefono:
        campos_derecha.append(("TELEFONO", telefono))
    if correo:
        campos_derecha.append(("CORREO", correo))
    if direccion:
        campos_derecha.append(("DIRECCION", direccion))
    if not campos_derecha:
        campos_derecha = [("ATENCION", datetime.now().strftime("%d/%m/%Y  ·  %H:%M"))]

    for i, (et, val) in enumerate(campos_izquierda):
        _insertar_campo(tabla.cell(0, 0), et, val, i == 0)
    for i, (et, val) in enumerate(campos_derecha):
        _insertar_campo(tabla.cell(0, 1), et, val, i == 0)

    _bordes_tabla(tabla, GRIS_BORDE_SUAVE_HEX, "4", bordes=("top", "left", "bottom", "right"))

    espaciador = doc.add_paragraph()
    espaciador.paragraph_format.space_after = Pt(2)


def generar_word_desde_nota_clinica(nota_clinica: str, tipo_documento: str, config: dict = None) -> bytes:
    global _archivos_temporales
    _archivos_temporales = []

    if config is None:
        config = {}

    try:
        doc = Document()
        _configurar_documento(doc)
        _agregar_encabezado(doc, config)
        _agregar_titulo(doc, tipo_documento)
        _agregar_datos_paciente(doc, config)
        _agregar_contenido_sin_campos_vacios(doc, nota_clinica)
        _agregar_firmas(doc, config)
        _agregar_pie(doc)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    finally:
        _limpiar_temporales()
